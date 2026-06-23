"""Generacion de reportes preliminares de EDA con modelos Gemini.

Este modulo concentra toda la logica de IA: lectura del diccionario/ODK,
serializacion de tablas ya calculadas por la app, construccion del prompt y
llamada al modelo. La interfaz vive en app.py (pestana "Informe IA").

Decision de diseno clave: la IA NO recibe la base cruda, solo recibe las
descriptivas y tablas cruzadas que la app ya calcula. Es mas barato, mas
preciso y mas seguro para datos de encuesta.
"""

from __future__ import annotations

import json
import re
from io import BytesIO
from typing import Any

import pandas as pd


# Modelos disponibles: solo los dos mas economicos recomendados.
# El nombre visible apunta al "model string" que entiende la API de Gemini.
AVAILABLE_MODELS: dict[str, str] = {
    "Gemini 2.5 Flash-Lite (mas economico)": "gemini-2.5-flash-lite",
    "Gemini 3.1 Flash-Lite": "gemini-3.1-flash-lite",
}


# Instruccion base que gobierna todo el reporte. El contexto del proyecto y las
# variables socioeconomicas son insumos obligatorios que se anexan en build_prompt.
SYSTEM_INSTRUCTION = """Actúa como un/a analista cuantitativo/a senior especializado/a en análisis exploratorio de datos de encuesta, monitoreo, evaluación y generación de reportes descriptivos.

Elabora en español un informe preliminar de análisis exploratorio de datos y relaciones entre variables usando únicamente las tablas agregadas provistas por la aplicación: descriptivas de variables continuas, frecuencias de variables categóricas, tablas cruzadas y diccionario de variables.

Integra el contexto del usuario para priorizar variables, grupos, dimensiones e hipótesis relevantes. No uses datos externos, no inventes cifras, no supongas información no contenida en las tablas y no hagas afirmaciones causales. Si algo no puede concluirse con la información disponible, indícalo explícitamente.

El reporte debe estar escrito en Markdown, con un estilo profesional, claro y narrativo. Prioriza una redacción narrativa en párrafos completos. Evita entregar el informe como una lista de bullets. Las cifras deben integrarse dentro de frases interpretativas y conectadas entre sí. Usa listas únicamente para próximos pasos o recomendaciones finales.

Usa transiciones entre ideas para explicar cómo se relacionan los hallazgos. No te limites a mencionar cifras: interpreta su relevancia para el contexto del estudio. Cuando presentes relaciones entre variables, explica el patrón observado, su posible importancia descriptiva y las precauciones necesarias para interpretarlo.

El informe debe incluir las siguientes secciones:

# Informe preliminar de análisis descriptivo

## 1. Resumen ejecutivo

Redacta entre dos y cuatro párrafos breves que sinteticen los principales hallazgos del análisis. Incluye la composición general de la muestra, los patrones descriptivos más importantes, las relaciones más relevantes entre variables y las principales advertencias metodológicas. No uses bullets en esta sección.

## 2. Contexto del análisis

Resume el objetivo del estudio o encuesta según el contexto aportado por el usuario. Explica cómo ese contexto orienta la lectura de los resultados. Si el contexto es insuficiente, indícalo de forma explícita y menciona cómo eso limita la interpretación.

## 3. Composición de la muestra

Describe en párrafos la composición de la muestra usando las variables disponibles. Incluye tamaño de muestra si está disponible y comenta la distribución por grupos relevantes como sexo, edad, territorio, zona, modalidad, institución, cohorte, grupo de intervención, exposición u otros segmentos disponibles.

## 4. Hallazgos descriptivos por dimensión

Organiza los hallazgos por temas o dimensiones, no necesariamente variable por variable. Agrupa variables relacionadas cuando sea posible. Para cada dimensión, redacta párrafos que combinen cifras concretas con interpretación. Evita presentar una lista mecánica de variables y valores.

## 5. Relaciones entre variables

Analiza los cruces disponibles buscando diferencias relevantes entre grupos, asociaciones descriptivas y patrones de segmentación. Prioriza los cruces entre variables directamente relacionadas con los objetivos del estudio y variables de resultado, participación, exposición, satisfacción, finalización, postulación, territorio, modalidad, zona, sexo, edad u otros grupos relevantes según el contexto.

Cada relación que describas DEBE estar cuantificada con las cifras de las tablas. No escribas que existe una diferencia, una asociación, una tendencia o un patrón sin acompañarlo del dato que lo respalda. Indica siempre los valores comparados de cada grupo y la magnitud de la diferencia (por ejemplo, el porcentaje de cada categoría, la diferencia en puntos porcentuales, las medias por grupo o los conteos), ya sea dentro de la frase o entre paréntesis. Ejemplo del nivel de sustento exigido: "la participación es mayor en zonas urbanas (62%) que en rurales (41%), una diferencia de 21 puntos porcentuales". Si una relación no puede cuantificarse con las tablas provistas, no la afirmes: menciónala como un cruce pendiente de analizar en los próximos pasos.

No te limites a describir frecuencias aisladas. Explica qué relación se observa, entre qué grupos aparece la diferencia, qué tan relevante parece ser y qué precauciones deben tenerse, siempre con las cifras de respaldo. Usa lenguaje descriptivo como "se observa", "los datos sugieren", "parece haber una diferencia" o "este patrón podría indicar". No hagas afirmaciones causales ni reportes significancia estadística si no está incluida en las tablas.

## 6. Calidad de los datos y limitaciones

Redacta una evaluación breve sobre posibles limitaciones del análisis: datos agregados, ausencia de base cruda, no respuesta, variables sin diccionario, cruces insuficientes, tamaños de celda pequeños, ausencia de pruebas estadísticas, imposibilidad de inferencia causal o cualquier otra restricción visible en las tablas.

## 7. Próximos pasos sugeridos

Presenta recomendaciones concretas para profundizar el análisis. En esta sección sí puedes usar una lista numerada, pero cada punto debe tener una breve explicación. Prioriza nuevos cruces, segmentaciones, visualizaciones, limpieza de variables, análisis de no respuesta, análisis de heterogeneidad, modelos estadísticos o validaciones adicionales según el contexto.

Reglas finales:

* Basa todo el reporte únicamente en las cifras provistas.
* Cada afirmación sobre la composición de la muestra, los hallazgos descriptivos o las relaciones entre variables debe ir acompañada del dato que la sustenta (cifra, porcentaje, diferencia en puntos porcentuales, media o n), expresado en el texto o entre paréntesis. No incluyas afirmaciones cualitativas sin su respaldo numérico; si no tienes el dato, no hagas la afirmación.
* No inventes datos, porcentajes, tamaños de muestra ni conclusiones.
* No uses conocimiento externo salvo para interpretar de forma general el tipo de análisis.
* No incluyas código.
* No incluyas fórmulas innecesarias.
* No muestres razonamiento paso a paso interno.
* Sé claro cuando una interpretación sea tentativa.
* Evita bullets en el resumen ejecutivo, contexto, composición de muestra, hallazgos descriptivos, relaciones entre variables y limitaciones.
* Usa bullets o listas solo cuando ayuden a organizar recomendaciones finales o advertencias muy puntuales.
* El resultado debe ser un borrador revisable por una persona antes de su uso.

Además del informe, sugiere gráficas que acompañen los hallazgos principales. No inventes variables y usa únicamente las variables seleccionadas por el usuario. Evita sugerir gráficos con demasiadas categorías o poco valor interpretativo.

Debes devolver la respuesta con este formato exacto:

<REPORT_MARKDOWN>
Aquí va únicamente el informe en Markdown.
</REPORT_MARKDOWN>

<CHART_SUGGESTIONS_JSON>
[
  {
    "chart_type": "bar",
    "title": "Título claro para la gráfica",
    "x": "variable_principal",
    "y": null,
    "color": "variable_opcional_o_null",
    "facet": "variable_opcional_o_null",
    "percent": true,
    "barmode": "group",
    "orientation": "v",
    "show_labels": true,
    "bins": null,
    "hole": 0.52,
    "show_trendline": false,
    "trendline_scope": "overall",
    "x_label": "Etiqueta eje X",
    "y_label": "Etiqueta eje Y",
    "legend_title": "Título de leyenda",
    "caption": "Texto breve que explique por qué esta gráfica acompaña el reporte."
  }
]
</CHART_SUGGESTIONS_JSON>

Reglas para el JSON:

* Devuelve JSON válido, sin comentarios ni Markdown dentro del bloque JSON. No uses cercas de código (```), comas finales ni texto fuera del arreglo.
* Devuelve EXACTAMENTE la cantidad de gráficas solicitada por el usuario siempre que existan suficientes variables adecuadas. Solo devuelve menos si no hay variables con valor interpretativo suficiente para alcanzar ese número; en ese caso, devuelve todas las que sí apliquen.
* Cierra siempre el arreglo con ] y el bloque con </CHART_SUGGESTIONS_JSON>.
* `chart_type` debe ser uno de: "bar", "histogram", "scatter", "pie".
* Para barras usa `x` como variable categórica principal. `color` y `facet` son opcionales.
* Para histogramas usa `x` como variable continua. `color` es opcional.
* Para dispersión usa `x` e `y` como variables continuas. `color` es opcional.
* Para pie usa `x` como variable categórica de proporción. `facet` es opcional.
* Si un campo no aplica, usa null, false o un valor razonable según el tipo.
* Cada sugerencia debe tener `title`, etiquetas y `caption`."""


# ---------------------------------------------------------------------------
# Lectura del diccionario / formulario ODK
# ---------------------------------------------------------------------------

def _read_uploaded_bytes(uploaded_file: Any) -> bytes:
    uploaded_file.seek(0)
    return uploaded_file.read()


def _find_label_column(columns: list[Any]) -> str | None:
    """Busca la columna de etiqueta. En ODK puede ser 'label' o 'label::Espanol (es)'."""
    lowered = {str(col).lower(): col for col in columns}
    if "label" in lowered:
        return lowered["label"]
    for col in columns:
        if str(col).lower().startswith("label"):
            return col
    return None


def _odk_mapping(survey: pd.DataFrame) -> tuple[Any, Any] | None:
    """Devuelve (columna_nombre, columna_etiqueta) si la hoja survey tiene estructura ODK."""
    lowered = {str(col).lower(): col for col in survey.columns}
    name_col = lowered.get("name")
    label_col = _find_label_column(list(survey.columns))
    if name_col is not None and label_col is not None:
        return name_col, label_col
    return None


def _inspect_generic_df(df: pd.DataFrame) -> dict[str, Any]:
    return {
        "kind": "generic",
        "df": df,
        "columns": [str(col) for col in df.columns],
        "name_col": None,
        "label_col": None,
    }


def inspect_dictionary_file(uploaded_file: Any) -> dict[str, Any]:
    """Lee el archivo de diccionario y detecta si es ODK o un diccionario generico.

    Retorna un dict con:
      - kind: "odk" | "generic"
      - df: DataFrame del diccionario (hoja survey para ODK)
      - columns: lista de columnas disponibles
      - name_col / label_col: columnas detectadas (solo en ODK; None en generico)
    """
    name = str(uploaded_file.name).lower()
    data = _read_uploaded_bytes(uploaded_file)
    if not data:
        raise ValueError("El archivo de diccionario esta vacio.")

    if name.endswith(".csv"):
        df = pd.read_csv(BytesIO(data), sep=None, engine="python")
        return _inspect_generic_df(df)

    if name.endswith((".xlsx", ".xls")):
        engine = "xlrd" if name.endswith(".xls") else "openpyxl"
        excel = pd.ExcelFile(BytesIO(data), engine=engine)
        sheets_lower = {str(sheet).lower(): sheet for sheet in excel.sheet_names}

        # Estructura ODK: hoja "survey" con columnas name/label.
        if "survey" in sheets_lower:
            survey = excel.parse(sheets_lower["survey"])
            mapping = _odk_mapping(survey)
            if mapping is not None:
                name_col, label_col = mapping
                return {
                    "kind": "odk",
                    "df": survey,
                    "columns": [str(col) for col in survey.columns],
                    "name_col": name_col,
                    "label_col": label_col,
                }
            # Hay hoja survey pero sin name/label -> tratar como generico.
            return _inspect_generic_df(survey)

        # Sin hoja survey: usar la primera hoja como diccionario generico.
        first = excel.parse(excel.sheet_names[0])
        return _inspect_generic_df(first)

    raise ValueError("Formato de diccionario no soportado. Usa CSV, XLSX o XLS.")


def build_variable_descriptions(
    dictionary_df: pd.DataFrame,
    name_col: Any,
    label_col: Any,
) -> dict[str, str]:
    """Construye el mapa nombre_variable -> descripcion a partir del diccionario."""
    descriptions: dict[str, str] = {}
    if dictionary_df is None or name_col is None or label_col is None:
        return descriptions
    if name_col not in dictionary_df.columns or label_col not in dictionary_df.columns:
        return descriptions

    for _, row in dictionary_df.iterrows():
        variable = row[name_col]
        label = row[label_col]
        if pd.isna(variable) or pd.isna(label):
            continue
        variable_text = str(variable).strip()
        label_text = str(label).strip()
        if variable_text and label_text:
            descriptions.setdefault(variable_text, label_text)
    return descriptions


# ---------------------------------------------------------------------------
# Serializacion de tablas a texto compacto (Markdown) para el prompt
# ---------------------------------------------------------------------------

def _format_cell(value: object, max_decimals: int) -> str:
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except (TypeError, ValueError):
        pass
    if isinstance(value, float):
        rounded = round(value, max_decimals)
        if float(rounded).is_integer():
            return str(int(rounded))
        return f"{rounded:.{max_decimals}f}".rstrip("0").rstrip(".")
    return str(value)


def _df_to_markdown(df: pd.DataFrame, *, index: bool, max_decimals: int) -> str:
    """Convierte un DataFrame a tabla Markdown sin depender de 'tabulate'."""
    if df is None or df.empty:
        return "(sin datos)"

    table = df.reset_index() if index else df.copy()
    headers = [str(col) for col in table.columns]
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for _, row in table.iterrows():
        cells = [_format_cell(value, max_decimals) for value in row.values]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def _descriptions_to_markdown(descriptions: dict[str, str]) -> str:
    if not descriptions:
        return "(no se proporciono diccionario de variables)"
    lines = ["| Variable | Descripcion |", "| --- | --- |"]
    for variable, label in descriptions.items():
        safe_label = label.replace("|", "/").replace("\n", " ")
        lines.append(f"| {variable} | {safe_label} |")
    return "\n".join(lines)


def build_data_context(
    continuous_table: pd.DataFrame,
    categorical_table: pd.DataFrame,
    crosstabs: list[tuple[str, str, pd.DataFrame]],
    table_type_label: str,
    *,
    max_decimals: int = 2,
) -> str:
    """Arma el bloque de datos que vera la IA a partir de las tablas calculadas."""
    parts: list[str] = []

    if continuous_table is not None and not continuous_table.empty:
        parts.append(
            "### Descriptivas de variables continuas\n"
            + _df_to_markdown(continuous_table, index=False, max_decimals=max_decimals)
        )
    else:
        parts.append("### Descriptivas de variables continuas\n(no se incluyeron variables continuas)")

    if categorical_table is not None and not categorical_table.empty:
        parts.append(
            "### Frecuencias de variables categoricas\n"
            + _df_to_markdown(categorical_table, index=False, max_decimals=max_decimals)
        )
    else:
        parts.append("### Frecuencias de variables categoricas\n(no se incluyeron variables categoricas)")

    if crosstabs:
        cross_parts: list[str] = []
        for main_var, disagg_var, table in crosstabs:
            cross_parts.append(
                f"#### {main_var} x {disagg_var} ({table_type_label})\n"
                + _df_to_markdown(table, index=True, max_decimals=max_decimals)
            )
        parts.append("### Tablas cruzadas\n" + "\n\n".join(cross_parts))
    else:
        parts.append("### Tablas cruzadas\n(no se incluyeron tablas cruzadas)")

    return "\n\n".join(parts)


def build_prompt(
    data_context: str,
    variable_descriptions: dict[str, str],
    user_context: str,
    *,
    dataset_name: str,
    n_rows: int,
    n_vars: int,
    socioeconomic_vars: list[str],
    selected_continuous_vars: list[str] | None = None,
    selected_categorical_vars: list[str] | None = None,
    max_chart_suggestions: int = 5,
    max_chart_categories: int = 15,
) -> str:
    """Construye el prompt final: instruccion base + insumos (contexto, diccionario, datos).

    El contexto del proyecto y las variables socioeconomicas son obligatorios; la
    validacion se hace en la interfaz antes de llamar a esta funcion.
    """
    dictionary_text = _descriptions_to_markdown(variable_descriptions)
    socio_text = ", ".join(socioeconomic_vars) if socioeconomic_vars else "(no se indicaron)"
    continuous_text = (
        ", ".join(selected_continuous_vars or [])
        if selected_continuous_vars
        else "(no se indicaron variables continuas)"
    )
    categorical_text = (
        ", ".join(selected_categorical_vars or [])
        if selected_categorical_vars
        else "(no se indicaron variables categóricas)"
    )

    return f"""{SYSTEM_INSTRUCTION}

---
# INSUMOS PARA EL INFORME

## 1. Contexto del proyecto, encuesta o estudio (aportado por el usuario)
{user_context.strip()}

## 2. Diccionario de variables
{dictionary_text}

## Informacion del dataset
- Nombre: {dataset_name}
- Numero de registros: {n_rows}
- Numero de variables: {n_vars}

## Variables socioeconomicas / de caracterizacion (obligatorias)
Las siguientes variables fueron marcadas como socioeconomicas o de caracterizacion y DEBEN usarse para describir la composicion de la muestra en la seccion correspondiente: {socio_text}

## Variables disponibles para sugerir graficas
- Continuas seleccionadas: {continuous_text}
- Categoricas seleccionadas: {categorical_text}
- Numero de graficas a sugerir: exactamente {max_chart_suggestions} si hay suficientes variables adecuadas; si no, todas las que tengan valor interpretativo.
- Usa unicamente las variables seleccionadas listadas arriba; no inventes ni uses otras.
- Evita repetir la misma variable principal en mas de una grafica salvo que aporte una vista claramente distinta.
- Evita graficas de barras o pie cuando la variable principal tenga mas de {max_chart_categories} categorias.

## 3 a 5. Datos calculados por la aplicacion (unica fuente de cifras permitida)
{data_context}
"""


# ---------------------------------------------------------------------------
# Llamada al modelo
# ---------------------------------------------------------------------------

def generate_report(
    api_key: str,
    model: str,
    prompt: str,
    *,
    max_output_tokens: int = 16384,
) -> str:
    """Llama a la API de Gemini y devuelve el texto del reporte.

    Lanza RuntimeError con un mensaje claro si falta la dependencia, si la
    respuesta viene vacia o si la API falla.
    """
    try:
        from google import genai
    except ImportError as exc:  # pragma: no cover - depende del entorno
        raise RuntimeError(
            "Falta la dependencia 'google-genai'. Instalala con: uv add google-genai"
        ) from exc

    client = genai.Client(api_key=api_key)

    # Se fija un limite alto de tokens de salida para que el informe junto con el
    # bloque de graficas no se trunquen (un JSON cortado deja sin sugerencias).
    config = None
    try:
        from google.genai import types
        config = types.GenerateContentConfig(max_output_tokens=max_output_tokens)
    except Exception:  # noqa: BLE001 - si la version del SDK no lo soporta
        config = None

    if config is not None:
        response = client.models.generate_content(model=model, contents=prompt, config=config)
    else:
        response = client.models.generate_content(model=model, contents=prompt)

    text = getattr(response, "text", None)
    if not text or not text.strip():
        raise RuntimeError(
            "El modelo no devolvio texto. Revisa la API key, el modelo seleccionado "
            "o los limites de uso de tu cuenta."
        )
    return text


# ---------------------------------------------------------------------------
# Separacion del reporte y las sugerencias de graficas
# ---------------------------------------------------------------------------

REPORT_START = "<REPORT_MARKDOWN>"
REPORT_END = "</REPORT_MARKDOWN>"
CHARTS_START = "<CHART_SUGGESTIONS_JSON>"
CHARTS_END = "</CHART_SUGGESTIONS_JSON>"


def _strip_code_fences(text: str) -> str:
    """Quita cercas de codigo tipo ```json ... ``` alrededor del bloque."""
    cleaned = text.strip()
    cleaned = re.sub(r"^```[a-zA-Z0-9]*\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    return cleaned.strip()


def _salvage_chart_objects(raw: str) -> list[dict[str, Any]]:
    """Rescata objetos {...} de un JSON malformado o truncado.

    Escanea el texto contando llaves (ignorando las que estan dentro de cadenas)
    y trata de parsear cada objeto de nivel superior. Asi, aunque el arreglo no
    cierre con ']' o el ultimo objeto venga cortado, se recuperan los completos.
    """
    objects: list[dict[str, Any]] = []
    depth = 0
    start: int | None = None
    in_string = False
    escape = False

    for index, char in enumerate(raw):
        if in_string:
            if escape:
                escape = False
            elif char == "\\":
                escape = True
            elif char == '"':
                in_string = False
            continue
        if char == '"':
            in_string = True
        elif char == "{":
            if depth == 0:
                start = index
            depth += 1
        elif char == "}":
            if depth > 0:
                depth -= 1
                if depth == 0 and start is not None:
                    candidate = raw[start : index + 1]
                    try:
                        parsed = json.loads(candidate)
                        if isinstance(parsed, dict):
                            objects.append(parsed)
                    except json.JSONDecodeError:
                        pass
                    start = None
    return objects


def split_report_response(response_text: str) -> tuple[str, list[dict[str, Any]]]:
    """Separa el Markdown del reporte y el JSON de graficas sugeridas.

    Es tolerante a respuestas imperfectas del modelo: acepta que falte la
    etiqueta de cierre, quita cercas de codigo y rescata objetos de un JSON
    parcial o truncado. Si no hay bloque de graficas, devuelve lista vacia.
    """
    text = response_text.strip()

    # --- Reporte en Markdown ---
    report = text
    if REPORT_START in text:
        after_start = text.split(REPORT_START, 1)[1]
        if REPORT_END in after_start:
            report = after_start.split(REPORT_END, 1)[0].strip()
        else:
            # Sin etiqueta de cierre: tomar hasta el bloque de graficas o el final.
            report = after_start.split(CHARTS_START, 1)[0].strip()
    elif CHARTS_START in text:
        report = text.split(CHARTS_START, 1)[0].strip()

    # --- Sugerencias de graficas ---
    suggestions: list[dict[str, Any]] = []
    if CHARTS_START in text:
        raw_json = text.split(CHARTS_START, 1)[1]
        if CHARTS_END in raw_json:
            raw_json = raw_json.split(CHARTS_END, 1)[0]
        raw_json = _strip_code_fences(raw_json)

        try:
            parsed = json.loads(raw_json)
            if isinstance(parsed, list):
                suggestions = [item for item in parsed if isinstance(item, dict)]
        except json.JSONDecodeError:
            suggestions = []

        # Si el parseo directo fallo o vino vacio, rescatar objetos completos.
        if not suggestions:
            suggestions = _salvage_chart_objects(raw_json)

    return report, suggestions


# ---------------------------------------------------------------------------
# Exportacion del reporte a Word (.docx) editable
# ---------------------------------------------------------------------------

def _is_table_separator(cells: list[str]) -> bool:
    """Detecta la fila separadora de una tabla Markdown (|---|---|)."""
    return all(cell and set(cell) <= set("-: ") for cell in cells)


def _add_runs_with_bold(paragraph, text: str) -> None:
    """Agrega texto a un parrafo interpretando **negritas** de Markdown."""
    for part in re.split(r"(\*\*.+?\*\*)", text):
        if len(part) > 4 and part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            paragraph.add_run(part)


def markdown_to_docx_bytes(
    markdown_text: str,
    figures: list[dict[str, Any]] | None = None,
) -> bytes:
    """Convierte el reporte en Markdown a un documento Word (.docx) editable.

    Soporta encabezados (#..####), listas con viñetas y numeradas, tablas
    Markdown y negritas en linea. Devuelve los bytes del .docx.
    """
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:  # pragma: no cover - depende del entorno
        raise RuntimeError(
            "Falta la dependencia 'python-docx'. Instalala con: uv add python-docx"
        ) from exc

    document = Document()
    lines = markdown_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    index = 0
    total = len(lines)

    while index < total:
        stripped = lines[index].strip()

        if not stripped:
            index += 1
            continue

        # Bloque de tabla Markdown
        if stripped.startswith("|") and stripped.endswith("|"):
            block: list[list[str]] = []
            while index < total and lines[index].strip().startswith("|"):
                cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
                block.append(cells)
                index += 1
            rows = [row for row in block if not _is_table_separator(row)]
            if rows:
                n_cols = max(len(row) for row in rows)
                table = document.add_table(rows=0, cols=n_cols)
                try:
                    table.style = "Light Grid Accent 1"
                except Exception:  # noqa: BLE001 - estilo opcional
                    pass
                for row in rows:
                    cells = table.add_row().cells
                    for col_index in range(n_cols):
                        cells[col_index].text = row[col_index] if col_index < len(row) else ""
            continue

        # Encabezados
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            document.add_heading(stripped[level:].strip(), level=min(max(level, 1), 4))
            index += 1
            continue

        # Listas con viñeta
        if stripped[:2] in ("- ", "* ") or stripped.startswith("• "):
            paragraph = document.add_paragraph(style="List Bullet")
            _add_runs_with_bold(paragraph, stripped[2:].strip())
            index += 1
            continue

        # Listas numeradas
        numbered = re.match(r"^\d+[.)]\s+(.*)", stripped)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            _add_runs_with_bold(paragraph, numbered.group(1))
            index += 1
            continue

        # Parrafo normal
        paragraph = document.add_paragraph()
        _add_runs_with_bold(paragraph, stripped)
        index += 1

    if figures:
        document.add_page_break()
        document.add_heading("Gráficas sugeridas", level=1)
        for index, figure in enumerate(figures, start=1):
            title = str(figure.get("title") or f"Gráfica sugerida {index}")
            caption = str(figure.get("caption") or "").strip()
            image_bytes = figure.get("png")
            document.add_heading(f"Figura {index}. {title}", level=2)
            if image_bytes:
                document.add_picture(BytesIO(image_bytes), width=Inches(6.5))
            if caption:
                paragraph = document.add_paragraph()
                _add_runs_with_bold(paragraph, caption)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
